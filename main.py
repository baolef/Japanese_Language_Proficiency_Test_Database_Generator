# !/usr/bin/python3
# -*- coding: utf-8 -*-
# @Project : Japanese Language Proficiency Test Database Generator
# @Time    : 2020/2/7 19:10
# @Author  : Fang Baole
# @Email   : fbl718@sjtu.edu.cn

import docx
import traceback
import functions as f
from word_edit import doc2docx, auto_num_check

if __name__ == '__main__':
    """The main function
    
    The function is the main part of the program. Its main goal is to distribute jobs.
    
    Args:
        None
    
    Returns:
        None
    """
    try:
        print('#####Japanese Language Proficiency Test Database Generator#####')
        print('###Author: Fang Baole   Version: 2.0   Update time: 2020/2/12###')
        filename = input('Please input the filename of the test paper: ')
        answer_name = input('Please input the filename of the answer: ')
        output_name = input(
            'Please input the filename of the output (Press Enter to use the default output filename): ')
        if not output_name:
            output_name = '日语题库.xlsx'

        print('Converting auto number to text in the file "' + filename + '"...')
        auto_num_check(filename)
        print('Conversion succeed.')
        print('Converting auto number to text in the file "' + answer_name + '"...')
        auto_num_check(answer_name)
        print('Conversion succeed.')

        try:
            file = docx.Document(filename)
        except:
            print('Converting "' + filename + '" into "' + filename + 'x" file...')
            filename = doc2docx(filename)
            print('Conversion succeed.')
            file = docx.Document(filename)

        try:
            answer = docx.Document(answer_name)
        except:
            print('Converting "' + answer_name + '" into "' + answer_name + 'x" file...')
            answer_name = doc2docx(answer_name)
            print('Conversion succeed.')
            answer = docx.Document(answer_name)

        ws, wb, y, m, l = f.add_sheet(file.paragraphs[0].text, output_name)
        mark = f.check_type(filename)
        answer_list = f.read_answer(answer, ws)
        f.read(file, answer_list[-5:], ws, wb, y, m, l, mark, output_name)
        f.replace(ws, ['_','＿'], '<under2></under2>')
        wb.save(output_name)
        print('The database is successfully output to "' + output_name + '".')
        input('Press any key to exit')
    except Exception as e:
        print(traceback.format_exc())
